"""
Document parsing service for PDF, DOCX, and TXT files.
"""

import io
from typing import Optional

class DocumentParser:
    """Parses various document formats to extract text."""
    
    def parse(self, content: bytes, content_type: str) -> str:
        """Parse document content based on MIME type."""
        parsers = {
            "application/pdf": self._parse_pdf,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document": self._parse_docx,
            "text/plain": self._parse_txt,
        }
        
        parser = parsers.get(content_type)
        if not parser:
            raise ValueError(f"Unsupported content type: {content_type}")
        
        return parser(content)
    
    def _parse_pdf(self, content: bytes) -> str:
        """Extract text from PDF using pdfplumber."""
        import pdfplumber
        
        text_parts = []
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
        
        return "\n\n".join(text_parts)
    
    def _parse_docx(self, content: bytes) -> str:
        """Extract text from DOCX using python-docx."""
        from docx import Document
        
        doc = Document(io.BytesIO(content))
        text_parts = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)
        
        return "\n\n".join(text_parts)
    
    def _parse_txt(self, content: bytes) -> str:
        """Parse plain text file."""
        try:
            return content.decode("utf-8")
        except UnicodeDecodeError:
            return content.decode("latin-1")
